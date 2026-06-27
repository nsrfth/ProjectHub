# -*- coding: utf-8 -*-
"""Generate the TaskHub vs PMIS-requirements gap-analysis report as a .docx."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = "C6EFCE"; AMBER = "FFEB9C"; RED = "FFC7CE"; HEAD = "1F4E79"; SUB = "DDEBF7"
FONT = "Arial"

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10)
# complex-script font (so Persian terms render)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
rfonts.set(qn('w:cs'), FONT); rfonts.set(qn('w:ascii'), FONT); rfonts.set(qn('w:hAnsi'), FONT)

for i, sz in [(1, 18), (2, 14), (3, 11)]:
    st = doc.styles[f"Heading {i}"]
    st.font.name = FONT; st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, color=None, size=9, white=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = Pt(size); run.bold = bold
    if white: run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # set cs font on run
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:cs'), FONT)

STATUS_FILL = {"Yes": GREEN, "Partial": AMBER, "Missing": RED}

def add_table(headers, rows, widths, status_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        set_cell(c, h, bold=True, white=True, size=9)
        shade(c, HEAD)
        c.width = Inches(widths[j])
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            set_cell(cells[j], val, size=9)
            cells[j].width = Inches(widths[j])
            if status_col is not None and j == status_col:
                key = val.split()[0] if val else ""
                if key in STATUS_FILL:
                    shade(cells[j], STATUS_FILL[key])
    doc.add_paragraph()
    return t

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def p(t, bold=False, italic=False, size=10):
    par = doc.add_paragraph()
    r = par.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size); r.font.name = FONT
    rpr = r._element.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:cs'), FONT); rpr.append(rf)
    return par
def bullet(t):
    par = doc.add_paragraph(style="List Bullet")
    r = par.add_run(t); r.font.size = Pt(10); r.font.name = FONT
    rpr = r._element.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:cs'), FONT); rpr.append(rf)

# ---- page number footer ----
def add_footer():
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("TaskHub Gap Analysis — Page ")
    run.font.size = Pt(8)
    fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
    para._p.append(fld1)

# ================= TITLE =================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("TaskHub — Gap Analysis Report")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1F,0x4E,0x79); r.font.name = FONT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Coverage of TaskHub against the «سامانه جامع مدیریت و کنترل پروژه» requirements")
r.font.size = Pt(12); r.italic = True; r.font.name = FONT
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Source documents: «سند نیازمندی‌های سامانه مدیریت پروژه» (PMO, بهمن ۱۴۰۴) + «شرح خدمات» (SoW)\n"
                 "Compared system: TaskHub v2.5.0  •  Report date: 2026-06-26")
r.font.size = Pt(9); r.font.name = FONT
doc.add_paragraph()
legend = doc.add_paragraph()
r = legend.add_run("Legend:  Yes = available built-in   •   Partial = needs customization/extension   •   "
                   "Missing = not present.   Effort: S ≈ 1–2 wks, M ≈ 3–6 wks, L ≈ 2–3+ months (one developer, indicative).")
r.font.size = Pt(9); r.italic = True; r.font.name = FONT
doc.add_page_break()

# ================= 1. EXECUTIVE SUMMARY =================
h1("1. Executive Summary")
p("TaskHub is a mature, self-hosted, bilingual (Farsi/English), Jalali-first project-management "
  "information system. It already covers the project-management core to a high standard: tasks, WBS, "
  "Gantt/CPM, baselines, Earned Value Management (EVM) with S-curve, cost control (CBS), resources, "
  "timesheets, risk, NCR, project records, correspondence (دبیرخانه), role-based access control, audit "
  "logging, and a documented REST API.")
p("Headline coverage:", bold=True)
bullet("SoW Requirement Matrix (15 mandatory / الزامی items): 7 fully met, 8 partial, 0 absent. Every "
       "mandatory line is at least partially satisfied — the partials are depth gaps, not missing foundations.")
bullet("11-package PMBOK requirements spec: ~60–65% functional coverage. Strong on Time, Cost/EVM, Risk, "
       "Resources, Stakeholders, Scope/WBS; weaker on contract administration depth, financial appraisal, "
       "engineering document control, HSE/Quality, executive group dashboards, and external integration.")
p("The seven decisive gaps for this tender:", bold=True)
for i, g in enumerate([
    "Contract-management depth — صورت‌وضعیت (progress invoices/IPC), claims, guarantees, AVL, tenders (مناقصه).",
    "Integration — financial system / ERP connector and Active Directory (SSO) sign-in.",
    "Financial appraisal — NPV/IRR, project cash-flow/liquidity, management & contingency reserves.",
    "Engineering document control — discipline-based register, MDL, transmittals, doc→schedule linkage.",
    "Quality & HSE — QC checklists, quality indicators, HSE incident/near-miss register.",
    "Executive group dashboards & report-builder — group tree-rollup with ﷼ values, geographic map, Word/Excel/PDF report generator.",
    "Microsoft Project import/sync (the spec assumes MSP performs core scheduling).",
], 1):
    bullet(f"{i}. {g}")
p("Verdict: TaskHub is a credible base platform that satisfies all mandatory items with light-to-moderate "
  "customization, but reaching the full PMBOK/EPC depth specified is a multi-phase build centred on the seven "
  "gaps above.", bold=True)

h2("Coverage scorecard")
add_table(
    ["Requirement set", "Yes", "Partial", "Missing", "Indicative coverage"],
    [["SoW — mandatory (الزامی), 15 items", "7", "8", "0", "100% reachable; 47% built-in"],
     ["SoW — desired (مطلوب), 4 items", "0", "1", "3", "~25%"],
     ["PMBOK spec — 11 packages", "3 strong", "6 partial", "2 weak", "~60–65%"]],
    [2.6, 0.7, 0.8, 0.8, 2.0])

# ================= 2. SCOPE & METHOD =================
h1("2. Scope & Method")
p("Two source documents were analyzed in full:")
bullet("«سند نیازمندی‌های سامانه مدیریت پروژه» — a PMBOK-based requirements specification (PMO unit, بهمن ۱۴۰۴) "
       "describing the PMIS as 11 software packages (§8.1–8.11) plus technical/non-functional requirements (§5).")
bullet("«شرح خدمات» (Statement of Work) — formal scope of services, a 19-line Requirement Matrix "
       "(الزامی/مطلوب), supplementary system requirements, execution and pricing rules.")
p("TaskHub capabilities are taken from the current codebase (backend & frontend, version 2.5.0). Status is "
  "assessed conservatively: a feature is marked “Yes” only when an end-to-end GUI + API exists today.")

# ================= 3. SOW MATRIX =================
doc.add_page_break()
h1("3. Part 1 — SoW Requirement Matrix Response")
p("The SoW asks each vendor to declare, per requirement, the coverage (موجود/سفارشی = built-in / "
  "customization), implementation time and cost. The table below provides coverage, status and an indicative "
  "effort; pricing is to be added in the financial proposal.")
sow_rows = [
 ["R1","Project definition + base info (CAPEX form, project charter / منشور)","الزامی","Partial","Customization","S–M","Projects + generic intake forms exist; add a charter/CAPEX template + approval."],
 ["R2","Work Breakdown Structure (WBS)","الزامی","Yes","Built-in","—","WBS tree, outline codes, % roll-up, create-child & reparent."],
 ["R3","Project schedule + Gantt chart","الزامی","Yes","Built-in","—","Gantt with CPM, lag/lead, milestones, baselines."],
 ["R4","Planned vs actual progress + schedule-variance report","الزامی","Yes","Built-in","—","Baselines & variance report; EVM SV/SPI."],
 ["R5","Task & action management","الزامی","Yes","Built-in","—","Status, RACI, approval gate, labels, custom fields, dependencies, recurrence."],
 ["R6","Management dashboards","الزامی","Partial","Customization","M","Configurable dashboards + Portfolio page; add group-rollup executive views."],
 ["R7","Progress dashboard + S-curve","الزامی","Yes","Built-in","—","EVM S-curve (PV/EV/AC)."],
 ["R8","Document management","الزامی","Partial","Customization","M","Records + attachments + correspondence; add engineering doc register."],
 ["R9","Contracts + progress invoices (صورت‌وضعیت)","الزامی","Partial","Customization","L","Contracts/POs/vendors exist; صورت‌وضعیت, claims, guarantees missing."],
 ["R10","User access levels","الزامی","Yes","Built-in","—","RBAC + custom roles + user groups + project grants."],
 ["R11","Activity audit log","الزامی","Yes","Built-in","—","Full create/edit/delete audit trail."],
 ["R12","Excel / PDF export","الزامی","Partial","Customization","S–M","CSV exports today; add PDF/Word/Excel report output."],
 ["R13","Approval workflows (گردش کار تأییدات)","الزامی","Partial","Customization","M","Fixed flows (task/timesheet/CR/expense); add configurable workflow engine."],
 ["R14","Automatic alerts & notifications","الزامی","Partial","Customization","S–M","Pull-based + due-date reminders; wire email/SMS delivery."],
 ["R15","API for system integration","الزامی","Yes","Built-in","—","OpenAPI/Swagger + scoped tokens + webhooks."],
 ["D1","Custom managerial reports","مطلوب","Partial","Customization","M","Dashboards exist; add ad-hoc report-builder."],
 ["D2","Project cash-flow management","مطلوب","Missing","Customization","M","Build cash-flow / liquidity timeline on the cost ledger."],
 ["D3","Tender management (مناقصه)","مطلوب","Missing","Customization","L","Build tender/bid module (docs, invitations, evaluation)."],
 ["D4","Advanced BI dashboards","مطلوب","Missing","Integration","M","Integrate a BI tool via the API, or build dedicated BI views."],
]
add_table(["#","Requirement","Type","Status","Coverage","Effort","Notes"],
          sow_rows, [0.4,2.3,0.7,0.7,0.9,0.5,2.6], status_col=3)

# ================= 4. PMBOK MODULES =================
doc.add_page_break()
h1("4. Part 2 — Detailed PMBOK Module Analysis (§8.1–8.11)")

def module(num, title_en, intro, headers, rows):
    h2(f"§8.{num} — {title_en}")
    if intro: p(intro, italic=True, size=9)
    add_table(headers, rows, [3.7, 0.8, 3.6], status_col=1)

module(1, "Proposed plans & integrated project management",
 "Central bank of proposals/approved plans/projects with financial indicators (NPV/IRR), charter, baselines, "
 "integrated change control across all sub-systems.",
 ["Required capability","Status","TaskHub notes"],
 [["Central database of approved projects","Yes","Projects stored centrally with full history."],
  ["Define new projects","Yes","Project creation with owner, team, budget, status."],
  ["Integrated change control across sub-systems","Partial","Change Requests + audit; no cross-module orchestration."],
  ["Keep change history on projects","Yes","Audit log + change-request records."],
  ["Lock baselines & approvals","Yes","Immutable baseline snapshots (capture/activate/compare)."],
  ["Proposal/idea stage + NPV/IRR appraisal","Missing","No pre-project proposal pipeline or financial appraisal."],
  ["FS (feasibility) preparation/update workflow","Missing","No feasibility-study workflow."],
  ["User report-builder","Partial","Configurable dashboards; no ad-hoc report-builder."]])

module(2, "Project cost management",
 "Budgeting on a Cost Breakdown Structure (CBS), cash-flow, EVM, change control on the cost baseline, "
 "management & contingency reserves.",
 ["Required capability","Status","TaskHub notes"],
 [["Monitor project costs","Yes","Cost summary: planned/committed/actual/remaining per currency."],
  ["CBS (cost breakdown structure)","Yes","Cost accounts tree (create/rename/delete)."],
  ["Compare planned vs actual + forecast","Yes","EVM CV/CPI + EAC (3 methods)."],
  ["Manage budget & cost-baseline changes","Partial","Budget lines + Change Requests; no formal cost-baseline change flow."],
  ["Prevent unapproved budget/cashflow/financing changes","Partial","Expense approve/reject; no locked financing approval chain."],
  ["Project cash-flow / liquidity","Missing","No time-phased cash-flow / liquidity view."],
  ["Management & contingency reserves (risk-based)","Missing","Not modeled."],
  ["Multi-currency roll-up + FX","Yes","FX rates + base-currency roll-up."],
  ["User report-builder","Partial","Dashboards only."]])

module(3, "Schedule management & progress control",
 "Main + sub-project scheduling, automatic progress, trend & delay analysis. The spec assumes Microsoft "
 "Project (MSP) performs core scheduling and the system integrates with it.",
 ["Required capability","Status","TaskHub notes"],
 [["Activity definition, scheduling, updates","Yes","Tasks + dependencies + CPM scheduling."],
  ["EPC planning based on WBS","Partial","WBS yes; explicit E/P/C phase tagging not modeled."],
  ["Record/control activity & work-package delays","Yes","Slip vs baseline + variance report."],
  ["Lock approved time baselines","Yes","Baseline capture/activate."],
  ["Graphical progress display","Yes","Gantt + planner charts + S-curve."],
  ["Time-baseline approval workflow","Partial","Baseline capture exists; no multi-step approval."],
  ["Progress-% receive & approve workflow","Partial","Task approval gate; no dedicated progress-approval step."],
  ["Microsoft Project import/sync","Missing","No MSP file import/round-trip."],
  ["User report-builder","Partial","Dashboards only."]])

module(4, "Project scope management (WBS)",
 "Detailed WBS, WBS coding & dictionary (PMBOK), scope lock, Change Request / Change Order.",
 ["Required capability","Status","TaskHub notes"],
 [["Define WBS with no depth limit","Yes","N-level WBS tree."],
  ["WBS coding","Yes","Auto outline codes (1, 1.1, 1.1.2 …)."],
  ["Lock approved scope (per approved FS)","Partial","Baseline lock yes; “per approved FS” not modeled."],
  ["WBS Dictionary (PMBOK)","Missing","No WBS dictionary entity."],
  ["Change Request","Yes","CR lifecycle Draft→Submitted→Approved→Applied."],
  ["Change Order (estimate→approve→issue) workflow","Partial","CR exists; full Change-Order issuance flow missing."],
  ["User report-builder","Partial","Dashboards only."]])

module(5, "Contract & procurement management",
 "Comprehensive contract bank: contracts, addenda, claims, alarms, guarantees, AVL, supplier performance, "
 "MTS progress, progress invoices (صورت‌وضعیت), tenders (مناقصه), punch lists. This is the widest gap.",
 ["Required capability","Status","TaskHub notes"],
 [["Contractor/consultant contracts","Yes","Contracts with status, value, vendor link."],
  ["Supplier master with profile/ranking/economic code","Partial","Vendors hold name/email/phone only."],
  ["Contract financial details (advance, guarantee, LC, insurance, deductions)","Missing","Not modeled."],
  ["Alarms on key contract dates","Missing","No scheduled alerting/email."],
  ["Addenda (الحاقیه) & handover MoMs","Missing","Not modeled."],
  ["Lock contractual obligations","Missing","Not modeled."],
  ["Contractual delay tracking","Missing","Not modeled."],
  ["Approved Vendor List (AVL)","Missing","Not modeled."],
  ["Supplier performance KPIs","Missing","Not modeled."],
  ["MTS supplier progress tracking","Missing","Not modeled."],
  ["Progress invoices (صورت‌وضعیت)","Missing","Not modeled — required."],
  ["Contract scheduling & progress control","Missing","Not modeled."],
  ["Claims workflow","Missing","Not modeled."],
  ["Tender (مناقصه) document prep & evaluation","Missing","Not modeled."],
  ["Punch-list (temporary/final handover)","Missing","Not modeled."],
  ["Purchase orders + auto commitment","Yes","PO→commitment auto-post on issue."]])

module(6, "Resource planning & control",
 "Resource bank, activity assignment, calendars, pay rates, attendance, performance evaluation.",
 ["Required capability","Status","TaskHub notes"],
 [["Assign activities to people/resources","Yes","Resource→task assignments with units & hours."],
  ["Resource catalog (human/equipment/material) + rates","Yes","Catalog + skills + cost/bill rate cards."],
  ["Working calendars","Partial","Capacity-calendar scaffold; team holidays."],
  ["Workload reports (tabular & graphical)","Partial","Tabular workload report; limited graphics."],
  ["Attendance (clock-in/out)","Missing","Not modeled."],
  ["Resource performance evaluation","Missing","Not modeled."],
  ["User report-builder","Partial","Dashboards only."]])

module(7, "Communications & engineering documents",
 "Unified document numbering/classification by discipline, MDL, status per document, link to schedule, "
 "version history, transmittals, daily/weekly/monthly reports.",
 ["Required capability","Status","TaskHub notes"],
 [["Engineering doc register by discipline","Partial","Records + custom fields can approximate; no native discipline register."],
  ["Status/stage per document","Partial","Record status workflow."],
  ["Link docs to schedule + auto-update progress","Missing","No document→schedule progress linkage."],
  ["Document distribution + access control","Partial","Project access + record visibility."],
  ["MDL (master document list) categorization","Missing","Not modeled."],
  ["Version history of engineering documents","Partial","Attachments without formal versioning."],
  ["Transmittals","Missing","Not modeled (correspondence is letter-based)."],
  ["Daily/weekly/monthly project reports","Partial","Records + correspondence; no structured periodic report."],
  ["User report-builder","Partial","Dashboards only."]])

module(8, "Stakeholder management",
 "Lessons learned, minutes of meeting (MoM), correspondence segregation, stakeholder register.",
 ["Required capability","Status","TaskHub notes"],
 [["Lessons learned register","Partial","Add as a custom record type (framework exists)."],
  ["Minutes of meeting (MoM) + follow-up","Yes","Built-in MoM record type + comment thread."],
  ["Project correspondence segregation","Yes","Correspondence (دبیرخانه) module + per-project register."],
  ["Stakeholder register","Yes","Built-in Stakeholder record type."]])

module(9, "Project risk management",
 "Identify/record risks, qualitative & quantitative analysis, prioritization, response plan, tracking & alerts.",
 ["Required capability","Status","TaskHub notes"],
 [["Record & assess risk (probability × impact)","Yes","5×5 matrix, auto score."],
  ["Response strategy & mitigation plan","Yes","Accept/Avoid/Mitigate/Transfer + plan + owner."],
  ["Track to final action / closure","Yes","Risk close workflow."],
  ["Risk ranking/prioritization","Yes","Score-based ranking."],
  ["Quantitative analysis (Monte Carlo)","Missing","Only qualitative scoring."],
  ["Response-planning approval workflow","Partial","No multi-step approval."],
  ["User report-builder","Partial","Dashboards only."]])

module(10, "Quality & HSE management",
 "Quality indicators, QC checklists, quality-management plan, NCR, CAPA; HSE incident/near-miss/event "
 "register, HSE plan, corrective/preventive actions.",
 ["Required capability","Status","TaskHub notes"],
 [["Non-conformance reports (NCR)","Yes","Severity, disposition, corrective-task link, close."],
  ["Quality indicators / KPIs","Missing","Not modeled."],
  ["QC checklists & plan","Missing","Not modeled."],
  ["CAPA (corrective/preventive)","Partial","NCR corrective task only."],
  ["HSE incident / near-miss / event register","Missing","Not modeled."],
  ["HSE plan & reports","Missing","Not modeled."],
  ["User report-builder","Partial","Dashboards only."]])

module(11, "Executive (group-level) dashboards",
 "Twelve management dashboards spanning the whole group/portfolio.",
 ["Dashboard (spec)","Status","TaskHub notes"],
 [["Overall E/P/C progress + financial + delivery at-a-glance","Partial","Portfolio + dashboards; not E/P/C-phased."],
  ["S-curve per project","Yes","EVM S-curve."],
  ["Delay Pareto (by cause)","Missing","No delay-cause analytics."],
  ["EVM comparison across projects","Partial","EVM per project; cross-project compare missing."],
  ["Portfolio financial comparison by C/P/E","Partial","Portfolio view; not C/P/E split."],
  ["Group tree-rollup with ﷼ values","Missing","No group financial tree."],
  ["Geographic map of projects","Missing","Not modeled."],
  ["Report-builder dashboards (Excel/PDF/Word)","Missing","CSV only."],
  ["Budget analysis (approved vs spent + forecast)","Partial","Cost summary + EVM."],
  ["Contract progress analysis + milestones","Missing","Not modeled."],
  ["Trend/forecast (budget needs over time)","Partial","EVM trend via snapshots."],
  ["Procedures & instructions bank","Partial","Records/docs could host it."]])

# ================= 5. NON-FUNCTIONAL =================
doc.add_page_break()
h1("5. Part 3 — Technical & Non-Functional Requirements")
add_table(["Requirement (§5 / الزامات تکمیلی)","Status","TaskHub notes"],
 [["Web-based, internal-network, future-extensible, scalable","Yes","React SPA + Fastify + PostgreSQL 16 + Redis + Caddy, Docker-deployed on-prem."],
  ["Bilingual UI Farsi/English with user switch","Yes","Full i18n EN+FA, runtime language switch."],
  ["Jalali (Shamsi) primary + Gregorian calendar","Yes","Jalali-first; Gregorian option; correct date conversion."],
  ["Unlimited projects + concurrent users","Yes","No structural limits."],
  ["Documented API (financial/ERP/office-automation/analytics)","Partial","OpenAPI + tokens + webhooks present; no prebuilt connectors."],
  ["Audit log (create/edit/delete history)","Yes","Built-in."],
  ["Flexible RBAC by role / project / data type","Yes","Custom roles, permissions, groups, project grants."],
  ["Responsive (PC / tablet / mobile)","Yes","Responsive web + PWA."],
  ["Reporting + dashboards + Excel/PDF export","Partial","Dashboards + CSV; PDF/Word/Excel report output to add."],
  ["Scalability / extensibility under load","Yes","Modular profiles; horizontally deployable."],
  ["Integration with financial system / ERP","Missing","No connector (API enables building one)."],
  ["Active Directory / SSO sign-in","Partial","SCIM user provisioning present; AD/LDAP/SAML login not wired."],
  ["Security: RBAC, activity log, backup, org auth","Yes","RBAC + audit + backups; org-auth (AD) partial."]],
 [3.4,0.8,3.9], status_col=1)

# ================= 6. CROSS-CUTTING =================
h1("6. Part 4 — Cross-cutting Gaps & Integration")
p("These themes cut across multiple modules and carry the most weight in evaluation:")
bullet("Integration: the spec repeatedly requires online links to the financial system, ERP, office automation "
       "and Active Directory. TaskHub exposes a documented API, webhooks and SCIM provisioning, but ships no "
       "ERP/financial connector and no AD/LDAP/SAML login. This is a hard requirement and a build item.")
bullet("Workflow engine: many packages require configurable گردش کار. TaskHub has fixed approval flows "
       "(task, timesheet, change-request, expense) rather than a general workflow designer.")
bullet("EPC phase model: explicit Engineering/Procurement/Construction phase tracking is not modeled.")
bullet("Email/SMS alerts: notifications are pull-based with due-date reminders; outbound email/SMS is not wired.")
bullet("Report-builder & exports: user-selectable fields with Word/Excel/PDF output is required in almost every "
       "package and is currently CSV-only.")
bullet("Microsoft Project: the schedule package assumes MSP for core scheduling; TaskHub uses its own CPM "
       "engine and offers no MSP import/round-trip.")

# ================= 7. GAP REGISTER =================
doc.add_page_break()
h1("7. Part 5 — Prioritized Gap Register")
p("P0 = mandatory for tender compliance · P1 = high value / explicitly required depth · P2 = desired.")
add_table(["Priority","Gap","Module","Effort"],
 [["P0","Active Directory / SSO sign-in","NFR / Security","M"],
  ["P0","Financial system / ERP integration connector","NFR / Integration","L"],
  ["P0","Contract administration: صورت‌وضعیت, claims, guarantees, addenda","§8.5","L"],
  ["P0","Configurable approval workflow engine","Cross-cutting","M–L"],
  ["P0","Word/Excel/PDF report-builder + exports","§8.11 / R12","M"],
  ["P0","Email/SMS alert delivery (contract dates, due reminders)","Cross-cutting","S–M"],
  ["P1","Financial appraisal: NPV/IRR, cash-flow, reserves","§8.1/§8.2","M–L"],
  ["P1","Engineering document register: disciplines, MDL, transmittals, versions","§8.7","L"],
  ["P1","Executive group dashboards: tree-rollup (﷼), EVM compare, delay Pareto","§8.11","M–L"],
  ["P1","Quality (QC checklists, indicators) + HSE incident register","§8.10","M–L"],
  ["P1","Tender (مناقصه) management","§8.5/D3","L"],
  ["P1","AVL + supplier performance KPIs","§8.5","M"],
  ["P1","Microsoft Project import/sync","§8.3","M"],
  ["P2","Project charter / CAPEX intake template","R1/§8.1","S–M"],
  ["P2","WBS Dictionary","§8.4","S"],
  ["P2","EPC phase model on tasks/WBS","§8.3","S–M"],
  ["P2","Resource attendance + performance evaluation","§8.6","M"],
  ["P2","Quantitative (Monte-Carlo) risk analysis","§8.9","M"],
  ["P2","Geographic project map","§8.11","S–M"]],
 [0.7,4.2,1.3,0.7])

# ================= 8. ROADMAP =================
h1("8. Part 6 — Recommended Roadmap")
p("Phase 1 — Tender-compliance baseline (≈ 8–12 wks):", bold=True)
for b in ["AD/SSO sign-in","Email/SMS alert delivery","Word/Excel/PDF report-builder",
          "Configurable approval workflow engine (v1)","Project charter / CAPEX intake template"]:
    bullet(b)
p("Phase 2 — Contract & financial depth (≈ 10–16 wks):", bold=True)
for b in ["صورت‌وضعیت (progress invoices), claims, guarantees, addenda","AVL + supplier KPIs",
          "Cash-flow / liquidity + management reserves","Financial system / ERP connector (first integration)"]:
    bullet(b)
p("Phase 3 — Engineering, quality & executive analytics (≈ 10–16 wks):", bold=True)
for b in ["Engineering document register (disciplines, MDL, transmittals, versions)",
          "QC checklists + HSE incident register","Executive group dashboards (tree-rollup, EVM compare, delay Pareto)",
          "Tender (مناقصه) management","MS Project import/sync"]:
    bullet(b)

# ================= 9. SUBMISSION CHECKLIST =================
h1("9. Part 7 — Proposal Submission Checklist (per SoW)")
p("The SoW requires the technical & financial proposal to include — independent of the product gaps above:")
for b in ["Detailed implementation Gantt (برنامه زمان‌بندی اجرا)",
          "Execution team introduction with CVs (معرفی تیم اجرایی – CV)",
          "Deployment methodology (متدولوژی استقرار)",
          "Deliverables list (لیست تحویل‌دادنی‌ها)",
          "Training plan (end users + system administrators) and full user documentation",
          "Support & SLA terms (پشتیبانی + توافق‌نامه سطح خدمات)",
          "Itemized pricing by module and by service (نصب، سفارشی‌سازی، آموزش، پشتیبانی سالانه)",
          "Documentation of architecture, database/scalability, infrastructure, security & integration",
          "IP & data ownership assigned to the client; technical & financial parts submitted separately; "
          "proposal validity ≥ 3 months; structure matching the SoW headings"]:
    bullet(b)

p("")
p("Prepared from the two supplied documents and a direct review of the TaskHub v2.5.0 codebase. Status "
  "assessments are conservative (GUI + API required for “Yes”). Effort figures are indicative for one "
  "developer and should be confirmed during detailed design.", italic=True, size=9)

add_footer()

out = os.path.join(os.path.expanduser("~"), "Desktop", "TaskHub_Gap_Analysis_Report.docx")
doc.save(out)
print("SAVED:", out)
